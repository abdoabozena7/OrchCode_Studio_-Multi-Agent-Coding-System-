from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\projects\Ai\OrchCode_Studio_(Multi-Agent-Coding-System)")
OUT = ROOT / "docs" / "reports" / "codex-level-assessment-hivo-studio-2026-06-22.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(92, 102, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths, vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = vertical_alignment
            set_cell_margins(cell)
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        tbl_grid.append(grid_col)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(24)
    title.font.color.rgb = INK
    title.font.bold = True
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(14)

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(16)
    h1.font.color.rgb = BLUE
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h2.font.size = Pt(13)
    h2.font.color.rgb = BLUE
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h3._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h3.font.size = Pt(12)
    h3.font.color.rgb = DARK_BLUE
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        st = styles[list_style]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.5)
        st.paragraph_format.first_line_indent = Inches(-0.25)
        st.paragraph_format.space_after = Pt(8)
        st.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.text = "Codex-Level Assessment | Hivo Studio / OrchCode Studio"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.text = "Generated assessment artifact - 2026-06-22"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED


def add_para(doc, text: str, style: str | None = None, bold_label: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_label:
        r = p.add_run(bold_label)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, title: str, body: str, fill: str = CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = INK
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.add_run(body)
    doc.add_paragraph()


def add_simple_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths, vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        r.bold = True
        r.font.color.rgb = INK
        r.font.size = Pt(10)
    repeat_header_row(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_margins(cells[idx])
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            run.font.size = Pt(10)
    for row in table.rows:
        prevent_row_split(row)
    set_table_widths(table, widths, vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP)
    doc.add_paragraph()
    return table


def repeat_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def add_score_table(doc):
    rows = [
        ("Core coding workflow", "15%", "Can write, understand, review, debug, and automate code in local projects.", "Patch proposal and Rust apply exist, but full end-to-end apply/verify is not consistently green.", "55%"),
        ("Code understanding and review", "15%", "Explains unfamiliar codebases and reviews bugs/edge cases.", "Fresh memory, evidence refs, and deep lanes exist; arbitrary deep project understanding is still incomplete.", "58%"),
        ("Provider reasoning truth", "12%", "Provider owns reasoning/prose while tools own facts and authority.", "ReasoningKernel v2 is strong, green in its focused suite, and avoids local assistant fallback.", "72%"),
        ("Safety and authority", "12%", "Approvals, sandboxing, workspace controls, and command safety.", "Rust path/patch/command gates are real; command policy is heuristic and not sandbox-grade.", "64%"),
        ("Memory and context", "10%", "Durable repo context and instructions reduce repeated discovery.", "SQLite memory is fresh and useful, but durable semantic understanding is not complete.", "66%"),
        ("Multi-agent parallelism", "8%", "Subagents and parallel work for exploration/testing/review.", "Swarm autopilot is architecturally rich, but provider-backed production paths and tests are not mature.", "50%"),
        ("Product UX surfaces", "10%", "CLI, app, IDE, worktrees, cloud, clear approvals, Git workflows.", "Desktop app exists with useful truth surfaces; no Codex-level CLI/IDE/cloud/worktree parity.", "42%"),
        ("Integrations and extensibility", "8%", "AGENTS.md, MCP, skills, plugins, automations, external connectors.", "AGENTS-style guidance exists; MCP/skills/plugins/automations are mostly absent as product features.", "28%"),
        ("Reliability and release readiness", "10%", "Green tests, stable lifecycle, replay, supportable failure modes.", "Typecheck, Rust, and ReasoningKernel pass; full runtime suite currently fails 156 tests.", "45%"),
    ]
    add_simple_table(
        doc,
        ["Capability", "Weight", "Codex-level expectation", "Hivo evidence", "Score"],
        rows,
        [1.25, 0.55, 1.75, 2.25, 0.7],
        header_fill=LIGHT_BLUE,
    )


def add_checklist(doc, title, items):
    doc.add_heading(title, level=2)
    rows = []
    for priority, item, why, work, acceptance in items:
        rows.append((priority, "[ ] " + item, why, work, acceptance))
    add_simple_table(
        doc,
        ["Pri.", "Modification", "Reason", "Work", "Done when"],
        rows,
        [0.45, 1.45, 1.35, 1.8, 1.45],
        header_fill=LIGHT_BLUE,
    )


def build_doc() -> None:
    doc = Document()
    style_document(doc)
    doc.core_properties.title = "Codex-Level Assessment for Hivo Studio / OrchCode Studio"
    doc.core_properties.author = "Codex"
    doc.core_properties.subject = "Assessment, parity score, and modification checklist"

    title = doc.add_paragraph(style="Title")
    title.add_run("Codex-Level Assessment for Hivo Studio / OrchCode Studio")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Pros and cons of Codex, estimated parity percentage, and detailed modification checklist")

    meta_rows = [
        ("Repository", r"D:\projects\Ai\OrchCode_Studio_(Multi-Agent-Coding-System)"),
        ("Assessment date", "2026-06-22"),
        ("Codex source basis", "Fresh OpenAI Codex manual fetch from developers.openai.com on 2026-06-22"),
        ("Repository evidence basis", "README, architecture docs, source inspection, memory status, typecheck, tests, and Rust validation"),
        ("Important note", "This is a practical maturity estimate, not a claim that Hivo can or should clone OpenAI Codex exactly."),
    ]
    add_simple_table(doc, ["Field", "Value"], meta_rows, [1.55, 4.95], header_fill=LIGHT_BLUE)

    add_callout(
        doc,
        "Bottom line",
        "Hivo Studio is architecturally ambitious and has several Codex-grade building blocks, especially ReasoningKernel v2, SQLite-backed memory, and Rust-owned patch/command authority. The whole product is not yet Codex-level: current verified parity is about 55% overall, with a reasonable uncertainty range of 52-60%. The fastest path upward is not a redesign; it is green validation, stronger provider recovery, durable event authority, and product-surface parity.",
    )

    doc.add_heading("Executive Summary", level=1)
    add_para(doc, "The project is no longer only a mock coding-agent shell. It has a real TypeScript runtime, a React/Tauri desktop, a Rust authority layer, SQLite-first project memory, provider-authored reasoning, safety gates, and internal swarm architecture. Those are strong foundations.")
    add_para(doc, "The gap against Codex is mostly product maturity and operational reliability: full runtime tests are failing, provider-unavailable paths still block major features, some docs still describe older mock/demo assumptions, session replay is not single-source authoritative for every lifecycle, command safety remains heuristic, and Codex-style extension surfaces such as MCP, skills, plugins, automations, cloud/worktree execution, and IDE integration are not yet comparable.")
    add_para(doc, "Recommended headline score: 55% of practical Codex-level parity today. Core architecture is closer to 70-75%; verified product readiness is closer to 40-45%.")

    add_simple_table(
        doc,
        ["Signal", "Result", "Meaning"],
        [
            ("Memory index", "Fresh, schema v2, 433 indexed files", "Repository context can be trusted for this assessment."),
            ("TypeScript typecheck", "Passed", "Protocol, runtime, and desktop TypeScript compile cleanly."),
            ("ReasoningKernel v2 suite", "42 passed / 0 failed", "The provider-truth and adaptive reasoning core is strong."),
            ("Rust/Tauri tests", "20 plus 18 passed / 0 failed", "Workspace, command, patch, and SQLite authority tests are healthy."),
            ("Full runtime suite", "503 passed / 156 failed", "End-to-end runtime product maturity is not yet release-grade."),
            ("Worktree state", "Already dirty before report", "The report avoids modifying source code and adds only a generated document artifact."),
        ],
        [1.65, 1.6, 3.25],
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("How To Interpret The Percentage", level=1)
    add_para(doc, "The percentage answers: if Codex is treated as the public product benchmark for a modern coding agent, how close is Hivo Studio today in implemented, validated, operator-usable capability?")
    add_para(doc, "The score deliberately separates architecture from shipped reliability. A system can have an impressive design and still score lower if the full suite fails, provider paths are brittle, or user-facing workflows are incomplete.")
    add_score_table(doc)
    add_callout(
        doc,
        "Score formula",
        "Weighted score from the table is about 55%. This should be treated as a maturity estimate, not a scientific benchmark. Hivo should not claim an 80% reasoning guarantee until the exact router/author/verifier/embedding profile has a passing gate-specific certification record from the sealed adaptive-reasoning eval process.",
    )

    doc.add_heading("Codex: Pros And Cons", level=1)
    doc.add_heading("Pros", level=2)
    add_bullets(doc, [
        "Mature coding-agent product surface: Codex can write code, understand unfamiliar codebases, review code, debug failures, and automate development tasks.",
        "Multiple operating surfaces: official docs describe CLI, IDE extension, app, and cloud-oriented workflows, which helps users choose the right surface for local, editor-attached, parallel, or hosted work.",
        "Good context discipline: Codex supports durable project guidance through AGENTS.md, plus task context, prompts, skills, MCP, plugins, and automations.",
        "Operator workflow: Codex emphasizes planning, approvals, sandboxing, command execution, diff review, and validation instead of treating model output as automatically trusted.",
        "Extensibility: skills, plugins, app connectors, MCP servers, and subagents give Codex a growing ecosystem instead of a single hardcoded agent loop.",
        "Parallelism and isolation: worktrees, cloud/local modes, and subagents let Codex split exploration or implementation without overloading the main thread.",
        "Product polish: Codex has a clearer user-facing lifecycle, documentation, and integration story than a project still stabilizing its own contracts.",
    ])
    doc.add_heading("Cons", level=2)
    add_bullets(doc, [
        "Closed product behavior: users can configure Codex, but cannot deeply rewrite its internal scheduler, reasoning kernel, or product roadmap.",
        "Account, plan, and rollout dependency: access and capabilities may depend on plan, workspace policy, platform, or current product rollout.",
        "Less suitable as a white-label orchestration platform: Codex is a coding-agent product, while Hivo is trying to become a customizable multi-agent factory.",
        "Cloud or connector usage can introduce governance questions for teams with strict data residency or self-hosting requirements.",
        "Deep customization requires learning several surfaces: AGENTS.md, config.toml, skills, plugins, MCP, hooks, and automations can be powerful but complex.",
        "Provider choice is not the same as Hivo's configurable provider layer. Hivo can be designed around local or OpenAI-compatible providers if that remains a product priority.",
    ])

    doc.add_heading("Where Hivo Is Strong", level=1)
    add_bullets(doc, [
        "Architecture is aligned with a serious coding-agent direction: the README describes a Tauri desktop, TypeScript runtime, SQLite memory, adaptive provider reasoning, swarm planning, approval gates, and Rust-owned patch/command authority.",
        "ReasoningKernel v2 is the best current parity asset. It requires provider-authored understanding, tool rounds, verification, repair, explicit budgets, and no local assistant message on operational provider failure.",
        "The Rust authority layer is credible. Tests passed for workspace guards, command policy, patch preflight, Rust git snapshots, and SQLite runtime event projections.",
        "SQLite memory is real and fresh. Memory status reports repo index, manifest, symbol index, file summaries, command inventory, project glossary, project intelligence, task history, run artifacts, campaign artifacts, and eval artifacts.",
        "The project is unusually honest about not claiming certification. It explicitly separates read_reasoning and action_reasoning and refuses an 80% guarantee without sealed holdout records.",
        "Swarm autopilot has detailed architecture for automatic staffing, executor caps, read-only ratios, scheduler traces, trial labs, and 300 logical-agent stress scenarios.",
        "The desktop UI already exposes provider telemetry, final response source, certification gates, evidence count, information gain, verifier verdict, restore truth, and patch state. That is stronger than a simple chat UI.",
    ])

    doc.add_heading("Where Hivo Is Weak Against Codex", level=1)
    add_bullets(doc, [
        "The full runtime test suite is failing. A product cannot be called Codex-level while its broad runtime validation is red.",
        "Provider-unavailable paths still create many failures. The system is now honest about provider failure, but it needs stronger recovery, test fixtures, and user-facing remediation.",
        "Mock/demo migration is incomplete. Production now forbids demo/mock sessions, but older tests and docs still expect mock behavior.",
        "End-to-end lifecycle authority is split. Rust applies patches, the frontend reports results, and the runtime reconciles state. This is better than direct model writes, but still not a single durable source of truth.",
        "Session replay is improving but still not uniformly event-authoritative. Snapshot restore remains a fallback and should not be marketed as full replay.",
        "Command policy is heuristic. Rust classifies and blocks obvious risks, but this is not equivalent to sandbox-grade containment or a full process supervisor.",
        "Codex extension parity is missing. Hivo does not yet have mature equivalents for skills, plugins, MCP, automations, cloud mode, worktree mode, IDE extension, or external app connectors.",
        "Frontend testing and end-to-end workflow tests are much thinner than runtime tests.",
        "The UI needs a more sober operator-console treatment. Decorative typography and access labels like Full Access can overstate trust if backend guarantees are not exact.",
    ])

    doc.add_heading("Codex Feature Gap Matrix", level=1)
    add_simple_table(
        doc,
        ["Codex capability", "Hivo status", "Gap"],
        [
            ("Write, understand, review, debug, automate", "Partial", "Core mechanics exist, but broad runtime tests fail and deep understanding is not fully general."),
            ("AGENTS.md durable guidance", "Strong", "Repo has AGENTS guidance and follows it; needs productized discovery/override UX if Hivo becomes a user tool."),
            ("CLI workflow", "Partial", "Many npm CLIs exist, but no polished Codex-like terminal UI."),
            ("Desktop app", "Partial", "Tauri operator console exists, but UI reliability and state clarity need work."),
            ("IDE extension", "Missing", "No comparable editor-attached workflow."),
            ("Worktrees", "Missing or minimal", "No Codex-like worktree isolation and thread management."),
            ("Cloud mode", "Missing", "No hosted remote execution surface."),
            ("MCP", "Missing as product capability", "Architecture can add it, but current product does not expose Codex-style MCP setup/use."),
            ("Skills", "Missing as product capability", "No reusable skill loader comparable to Codex skills."),
            ("Plugins/connectors", "Missing", "No installable plugin/app ecosystem comparable to Codex."),
            ("Automations", "Missing or early", "Campaigns exist, but not Codex-style scheduled thread automations."),
            ("Subagents", "Partial", "Internal swarm exists, but provider-backed production subagent workflows are not fully mature."),
            ("Approvals and sandboxing", "Partial", "Approval gates and Rust authority are real; sandbox/process isolation is weaker."),
        ],
        [1.75, 1.3, 3.45],
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("Validation Findings", level=1)
    add_para(doc, "Commands run during this assessment:")
    add_simple_table(
        doc,
        ["Command", "Result", "Interpretation"],
        [
            ("npm run memory:index-status", "Passed; fresh index; 433 indexed files", "Context-sensitive assessment may rely on current repository memory."),
            ("npm run memory:status", "Passed", "SQLite memory is db_first and includes key index artifacts."),
            ("npm run typecheck", "Passed", "Strict TypeScript health is good across workspaces."),
            ("npm run test:reasoning-v2", "Passed; 42/42", "Adaptive reasoning and provider-truth kernel are in good shape."),
            ("cargo test --manifest-path apps/desktop/src-tauri/Cargo.toml", "Passed", "Rust authority layer is healthy in tested areas."),
            ("npm run test", "Failed; 503 passed, 156 failed", "Full runtime product validation is not currently green."),
        ],
        [2.35, 1.55, 2.6],
        header_fill=LIGHT_BLUE,
    )
    add_para(doc, "Visible failure categories in the full runtime suite:")
    add_bullets(doc, [
        "Tests still creating demo/mock sessions now fail with provider_mock_forbidden.",
        "Provider-backed explain and swarm tests fail when the provider is unavailable or marked not used.",
        "Some plan-only orchestration tests expect succeeded while current behavior returns blocked.",
        "Several UniversalProjectQuestionEngine tests fail because project_explain_provider_failed is now terminal rather than locally synthesized.",
    ])

    doc.add_heading("Detailed Modification Checklist", level=1)
    add_checklist(doc, "P0: Trust And Correctness Blockers", [
        ("P0", "Make the full runtime suite green", "A Codex-level product cannot have 156 failing runtime tests.", "Update mock-era tests to use private scripted real_provider fixtures, fix provider-backed test setup, and align plan-only expectations with current blocked semantics.", "npm run test passes with 0 failures."),
        ("P0", "Finish mock/demo migration", "Production now forbids mock sessions, but docs and tests still refer to mock runtime paths.", "Remove or quarantine runtime-selectable mock assumptions; keep scripted providers only inside tests; update module docs and smoke names.", "Production search finds no runtime-selectable mock path outside tests or history."),
        ("P0", "Create one durable event authority", "Patch/apply/verify truth is split across runtime, frontend, Rust, and SQLite projections.", "Standardize event names, make Rust apply/command results durable, and let runtime restore from authoritative events instead of frontend mirroring.", "Crash/restart test reconstructs patch approval, apply state, command results, and verification state from persisted events."),
        ("P0", "Fix provider-unavailable recovery", "Provider failure is honest but currently too brittle for deep project questions and tests.", "Separate expected provider-failure tests from real-provider happy paths; add actionable configuration errors; add bounded provider-guided repair before final cannot-answer.", "Provider outage produces a clear terminal state and no local answer; configured scripted provider tests pass."),
        ("P0", "Secure provider secret storage", "OpenAI-compatible providers need real secret handling before broader use.", "Add OS keychain or secure credential storage, avoid raw API key persistence, and show provider config health in UI.", "Provider config can survive restart without exposing secrets and without marking cloud providers invalid only because secrets are absent."),
        ("P0", "Do not claim the 80% reasoning guarantee", "Certification requires exact profile/gate evidence, not unit tests.", "Create sealed read and action holdout corpora, run eval:adaptive-reasoning, register only passing reports.", "Certification registry contains a valid passing report for the exact model profile and gate."),
    ])
    add_checklist(doc, "P1: Codex-Parity Product Work", [
        ("P1", "Build a true ProjectUnderstandingKernel", "Deep project questions still need relationship-aware understanding, not selected-snippet drafting.", "Implement provider decomposition, graph expansion, claim ledger, evidence validation, and validation-as-repair.", "Novel cross-file questions answer with current citations or explicit unknowns after bounded repair."),
        ("P1", "Replace heuristic command safety with structured policy", "Heuristics are useful but not sandbox-grade.", "Design command DSL/allowlist, adversarial parser tests, approval provenance, and a process supervisor for background jobs.", "Command tests cover shell indirection, wrappers, encoded/network behavior, background lifecycle, and cleanup."),
        ("P1", "Harden patch lifecycle end to end", "Rust apply is strong, but completion and reconciliation still cross layers.", "Make Rust-to-runtime completion durable, rerun validation through Rust, and reconcile after reconnect.", "E2E test covers propose -> approve -> Rust apply -> validate -> final verified report, plus frontend crash after apply."),
        ("P1", "Productize skills and reusable workflows", "Codex gains reliability from skills; Hivo currently relies on repo-specific architecture.", "Define Hivo skill format or adapt existing Agent Skills, loader, trigger rules, references, and script execution safety.", "A skill can be installed, selected, read progressively, and used in a run with an auditable trace."),
        ("P1", "Add MCP integration", "Codex uses MCP for external tools/context; Hivo lacks equivalent connector infrastructure.", "Add MCP server config, auth model, tool allowlists, trace logging, and UI for tool availability.", "A read-only docs MCP and one local tool MCP can be configured and used with audit logs."),
        ("P1", "Add worktree isolation", "Codex app supports worktrees for isolated changes; Hivo needs safer parallel execution.", "Implement branch/worktree creation, per-thread workspace roots, cleanup, and merge/review workflow.", "Two independent tasks can run in separate worktrees without modifying the base workspace."),
        ("P1", "Make provider-backed swarm production-safe", "Internal swarm has strong architecture but still includes mock/test-oriented paths.", "Require provider-backed read-only workers for real runs, remove auto mock fallback, durably track leases/retries/conflicts.", "A broad read-only scan completes with provider-backed workers, executor cap 0, and reviewed evidence inventory."),
        ("P1", "Redesign operator UX around trust states", "The UI must show exactly what is proposed, approved, applied, verified, failed, or unrecoverable.", "Replace decorative typography, refine access labels, elevate lifecycle banner, and make restore/apply/verification status explicit.", "User can identify next safe action within 5 seconds in manual QA scenarios."),
    ])
    add_checklist(doc, "P2: Scale, Polish, And Adoption", [
        ("P2", "Add frontend and end-to-end tests", "Runtime tests are strong but seam coverage is thin.", "Use Playwright or component tests for approval, restore, disconnect, provider failure, and patch review UI.", "CI includes UI state-machine tests and at least one full local runtime-to-Rust apply flow."),
        ("P2", "Introduce Codex-style automations", "Campaigns exist, but scheduled monitors and thread automations are not productized.", "Add scheduler, reminder/monitor definitions, execution logs, and archive/notification behavior.", "A recurring read-only report automation can run, persist evidence, and notify the operator."),
        ("P2", "Clean stale documentation", "Some docs still describe old Module 1/2/3 assumptions.", "Update security-model, module status docs, deep-dive references, and release notes to current provider/mock/replay truth.", "Docs no longer contradict production mock forbiddance or current Rust patch application."),
        ("P2", "Create release health dashboard", "Operators need one place to see provider, memory, tests, git, runtime, and Rust authority health.", "Add UI and CLI health commands that run safe checks and summarize blockers.", "A single command/report gives green/yellow/red status with remediation links."),
        ("P2", "Benchmark against real tasks", "Unit tests do not prove product usefulness.", "Create a held-out task corpus covering bug fixes, refactors, tests, docs, UI changes, and multi-file understanding.", "Monthly report tracks task success, unsupported claims, validation success, latency, and operator interventions."),
    ])

    doc.add_heading("Suggested Roadmap", level=1)
    add_simple_table(
        doc,
        ["Horizon", "Goal", "Expected parity movement"],
        [
            ("0-30 days", "Make validation green, update mock-era tests/docs, fix provider test setup, and clarify UI trust labels.", "55% -> 62-65%"),
            ("30-60 days", "Unify durable event authority, improve session replay, harden patch/apply/verify E2E, and add frontend tests.", "62-65% -> 68-72%"),
            ("60-90 days", "Implement ProjectUnderstandingKernel, structured command policy, provider-backed read-only swarm, and worktree isolation.", "68-72% -> 75-80%"),
            ("90+ days", "Add skills, MCP, plugins/connectors, automations, IDE/CLI polish, and sealed certification gates.", "75-80% -> 82-88% if validation remains green."),
        ],
        [1.0, 3.8, 1.7],
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("What Not To Do", level=1)
    add_bullets(doc, [
        "Do not chase a cosmetic Codex clone before trust and validation are green.",
        "Do not add local semantic fallback answers to make demos look better.",
        "Do not widen command or patch authority to compensate for missing lifecycle repair.",
        "Do not market 300 agents as the normal user experience; keep it a maximum internal capacity.",
        "Do not claim Codex-level reasoning from unit tests or scripted smoke runs.",
        "Do not solve deep project understanding with more hardcoded concept rules.",
    ])

    doc.add_heading("Evidence Appendix", level=1)
    add_para(doc, "Official Codex evidence used:")
    add_bullets(doc, [
        "OpenAI Codex manual, fetched on 2026-06-22 from developers.openai.com/codex/codex-manual.md.",
        "Manual sections consulted: Codex overview, best practices, Codex app features, CLI features, AGENTS.md customization, skills, MCP, subagents, and plugins.",
    ])
    add_para(doc, "Repository evidence used:")
    add_bullets(doc, [
        "README.md and docs/architecture.md for system architecture.",
        "docs/agent-contracts.md for ReasoningKernel and provider-truth contracts.",
        "docs/adaptive-reasoning-certification.md for certification gate requirements.",
        "docs/release-candidate-status.md and docs/reliability-audit.md for known limitations.",
        "docs/project-understanding-general-intelligence-diagnosis.md for deep understanding gaps.",
        "apps/agent-runtime/src/runtime/ReasoningKernel.ts and related tests for provider-authored reasoning.",
        "apps/desktop/src-tauri/src/commands/patch.rs and services/patch.rs for Rust patch authority.",
        "apps/desktop/src-tauri/src/services/terminal.rs and command_policy.rs for command authority.",
        "apps/agent-runtime/src/memory and .agent_memory status output for SQLite memory.",
        "apps/agent-runtime/src/orchestration for swarm, trials, team sub-planning, and validation gates.",
    ])

    doc.add_heading("Final Recommendation", level=1)
    add_para(doc, "Treat Hivo as a promising orchestration-first coding-agent platform, not yet as a Codex replacement. The architecture is good enough to continue. The immediate focus should be product hardening: green tests, durable truth, provider recovery, and a stricter operator UX. After that, the unique Hivo thesis becomes compelling: a local, auditable, multi-agent coding factory with provider-configurable intelligence and deterministic authority boundaries.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
